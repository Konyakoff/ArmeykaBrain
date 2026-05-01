"""Export article (markdown) to HTML / DOCX / PDF with consistent styling."""

from __future__ import annotations

import io
import re
import logging
from typing import Literal

import markdown
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

logger = logging.getLogger("export")

ExportFormat = Literal["html", "docx", "pdf"]

_BRAND_ORANGE = "#F47920"

_CSS = """\
@import url('https://fonts.googleapis.com/css2?family=Montserrat:wght@400;500;600;700;800;900&family=Onest:wght@700;800&display=swap');

body {
    font-family: 'Montserrat', 'Segoe UI', Arial, sans-serif;
    font-size: 14px;
    line-height: 1.7;
    color: #374151;
    max-width: 820px;
    margin: 40px auto;
    padding: 0 32px;
    background: #fff;
}
h1, h2, h3 {
    font-family: 'Onest', 'Montserrat', sans-serif;
    color: #333;
    font-weight: 800;
    margin-top: 1.75em;
    margin-bottom: 0.75em;
    line-height: 1.3;
}
h1 { font-size: 1.5em; }
h2 { font-size: 1.25em; }
h3 { font-size: 1.1em; }
p  { margin-bottom: 1em; }
ul { list-style: none; margin-left: 0; margin-bottom: 1em; padding-left: 0; }
ul li {
    position: relative;
    padding-left: 1.5em;
    margin-bottom: 0.5em;
}
ul li::before {
    content: '\\2022';
    color: %(orange)s;
    font-weight: bold;
    position: absolute;
    left: 0;
    font-size: 1.2em;
    line-height: 1;
}
ol { list-style-type: decimal; margin-left: 1.5em; margin-bottom: 1em; }
li { margin-bottom: 0.35em; }
strong { font-weight: 700; color: #333; }
em { font-style: italic; }
blockquote {
    border: 2px solid #f0dfd1;
    border-left: 4px solid %(orange)s;
    background: #fcf6f2;
    padding: 1.25em 1.5em;
    color: #555;
    border-radius: 12px;
    margin: 1.5em 0;
    font-style: normal;
}
pre {
    border: 2px solid #f0dfd1;
    border-left: 4px solid %(orange)s;
    background: #fcf6f2;
    padding: 1.25em 1.5em;
    color: #555;
    border-radius: 12px;
    margin: 1.75em 0;
    white-space: pre-wrap;
    word-wrap: break-word;
    line-height: 1.6;
}
code {
    background: #fcf6f2;
    color: #d46b08;
    padding: 0.2em 0.4em;
    border-radius: 0.25rem;
    font-weight: 600;
}
pre code { background: transparent; color: inherit; padding: 0; }
a { color: %(orange)s; text-decoration: underline; font-weight: 500; }
table { border-collapse: collapse; width: 100%%; margin: 1em 0; }
th, td { border: 1px solid #e5e7eb; padding: 8px 12px; text-align: left; font-size: 13px; }
th { background: #f8fafc; font-weight: 700; color: #333; }
""" % {"orange": _BRAND_ORANGE}


def _md_to_html(md_text: str) -> str:
    return markdown.markdown(
        md_text,
        extensions=["tables", "fenced_code", "nl2br", "sane_lists"],
    )


def _full_html(title: str, md_text: str) -> str:
    body = _md_to_html(md_text)
    return (
        "<!DOCTYPE html>\n"
        '<html lang="ru">\n<head>\n'
        '<meta charset="UTF-8">\n'
        '<meta name="viewport" content="width=device-width,initial-scale=1">\n'
        f"<title>{title}</title>\n"
        f"<style>{_CSS}</style>\n"
        "</head>\n<body>\n"
        f"{body}\n"
        "</body>\n</html>"
    )


def export_html(title: str, md_text: str) -> bytes:
    return _full_html(title, md_text).encode("utf-8")


def export_pdf(title: str, md_text: str) -> bytes:
    from weasyprint import HTML as WPHTML
    html_str = _full_html(title, md_text)
    return WPHTML(string=html_str).write_pdf()


_ORANGE_RGB = RGBColor(0xF4, 0x79, 0x20)
_DARK_RGB = RGBColor(0x33, 0x33, 0x33)
_GRAY_RGB = RGBColor(0x55, 0x55, 0x55)
_WHITE_RGB = RGBColor(0xFF, 0xFF, 0xFF)


def _set_run_props(run, *, bold=False, italic=False, size=None, color=None, font_name=None):
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if size:
        run.font.size = size
    if color:
        run.font.color.rgb = color
    if font_name:
        run.font.name = font_name


def _strip_md_link(text: str) -> str:
    """[text](url) -> text"""
    return re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)


def _process_inline(paragraph, text: str):
    """Render markdown inline formatting (**bold**, *italic*, `code`) into docx runs."""
    text = _strip_md_link(text)
    tokens = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)', text)
    for tok in tokens:
        if not tok:
            continue
        if tok.startswith('**') and tok.endswith('**'):
            run = paragraph.add_run(tok[2:-2])
            _set_run_props(run, bold=True, color=_DARK_RGB, size=Pt(11))
        elif tok.startswith('*') and tok.endswith('*'):
            run = paragraph.add_run(tok[1:-1])
            _set_run_props(run, italic=True, size=Pt(11))
        elif tok.startswith('`') and tok.endswith('`'):
            run = paragraph.add_run(tok[1:-1])
            _set_run_props(run, bold=True, color=RGBColor(0xD4, 0x6B, 0x08), size=Pt(10))
        else:
            run = paragraph.add_run(tok)
            _set_run_props(run, size=Pt(11), color=_GRAY_RGB)


def export_docx(title: str, md_text: str) -> bytes:
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.font.color.rgb = _GRAY_RGB
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.5

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    lines = md_text.split('\n')
    in_code_block = False
    code_lines: list[str] = []
    in_list = False

    for line in lines:
        stripped = line.strip()

        if stripped.startswith('```'):
            if in_code_block:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(8)
                p.paragraph_format.space_after = Pt(8)
                run = p.add_run('\n'.join(code_lines))
                _set_run_props(run, size=Pt(10), color=_GRAY_RGB, font_name='Consolas')
                shading = p.paragraph_format.element
                pPr = shading.get_or_add_pPr()
                shd = pPr.makeelement(qn('w:shd'), {
                    qn('w:fill'): 'FCF6F2',
                    qn('w:val'): 'clear',
                })
                pPr.append(shd)
                code_lines = []
            in_code_block = not in_code_block
            continue

        if in_code_block:
            code_lines.append(stripped)
            continue

        if not stripped:
            in_list = False
            continue

        heading_match = re.match(r'^(#{1,3})\s+(.*)', stripped)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2)
            p = doc.add_heading(level=level)
            run = p.add_run(text)
            _set_run_props(
                run,
                bold=True,
                color=_DARK_RGB,
                font_name='Calibri',
                size=Pt(16 - level * 2),
            )
            in_list = False
            continue

        if stripped.startswith('>'):
            text = stripped.lstrip('> ')
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            pPr = p.paragraph_format.element.get_or_add_pPr()
            shd = pPr.makeelement(qn('w:shd'), {
                qn('w:fill'): 'FCF6F2',
                qn('w:val'): 'clear',
            })
            pPr.append(shd)
            border_el = pPr.makeelement(qn('w:pBdr'), {})
            left_border = border_el.makeelement(qn('w:left'), {
                qn('w:val'): 'single',
                qn('w:sz'): '12',
                qn('w:space'): '4',
                qn('w:color'): 'F47920',
            })
            border_el.append(left_border)
            pPr.append(border_el)
            _process_inline(p, text)
            continue

        bullet_match = re.match(r'^[-*+]\s+(.*)', stripped)
        if bullet_match:
            text = bullet_match.group(1)
            p = doc.add_paragraph(style='List Bullet')
            _process_inline(p, text)
            in_list = True
            continue

        num_match = re.match(r'^\d+[.)]\s+(.*)', stripped)
        if num_match:
            text = num_match.group(1)
            p = doc.add_paragraph(style='List Number')
            _process_inline(p, text)
            in_list = True
            continue

        p = doc.add_paragraph()
        _process_inline(p, stripped)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def export_article(title: str, md_text: str, fmt: ExportFormat) -> tuple[bytes, str, str]:
    """Return (bytes, content_type, file_extension)."""
    if fmt == "html":
        return export_html(title, md_text), "text/html; charset=utf-8", ".html"
    if fmt == "docx":
        ct = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        return export_docx(title, md_text), ct, ".docx"
    if fmt == "pdf":
        return export_pdf(title, md_text), "application/pdf", ".pdf"
    raise ValueError(f"Unknown format: {fmt}")
